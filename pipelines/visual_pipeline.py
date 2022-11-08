import atexit
import functools
import io
import json
import random
import re
from collections.abc import Callable
from concurrent.futures import ThreadPoolExecutor, as_completed
from hashlib import sha256
from logging import ERROR, getLogger
from operator import or_
from pathlib import Path
from typing import Literal

import docx
import numpy as np
import PIL.Image as Image
import pyparsing as pp
import requests
import tqdm
from lxml import etree
from nltk import word_tokenize
from pyparsing import common as ppc
from pyquery import PyQuery as pq

from pipelines.utils import register

logger = getLogger(__name__)
temp_dir = Path("temp")
temp_dir.mkdir(exist_ok=True)


def rm_dir(path: Path):
    """Recursively remove the given directory."""

    for child in path.iterdir():
        if child.is_dir():
            rm_dir(child)
        else:
            child.unlink()
    path.rmdir()


atexit.register(functools.partial(rm_dir, temp_dir))

ua = "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.114 Safari/537.36"


def get(url, *args, **kwargs):
    """Get the given url and return the response."""
    return requests.get(url, *args, headers={"User-Agent": ua}, **kwargs)


def get_image(description: str) -> Image.Image:
    """Get an image from Google Images that is relevant to the given description."""
    from icrawler.builtin import GoogleImageCrawler

    temp = temp_dir / sha256(description.encode("utf-8")).hexdigest()
    google_crawler = GoogleImageCrawler(
        feeder_threads=1,
        parser_threads=1,
        downloader_threads=1,
        storage={"root_dir": temp},
    )
    google_crawler.set_logger(ERROR)
    google_crawler.crawl(keyword=description, max_num=4)

    return Image.open(random.choice(list(temp.iterdir())))


def get_icon(
    word: str,
    color: Literal["black", "gradient", "color"] = "color",
    shape: Literal["all", "hand-drawn", "lineal-color", "fill", "outline"] = "fill",
) -> Image.Image:

    html = etree.HTML(
        get(
            "https://www.flaticon.com/search",
            params={
                "word": word,
                "type": "icon",
                "license": "",
                "color": color,
                "shape": shape,
            },
        ).text
    )
    icon_page_url = random.choice(html.xpath('//a[contains(@href, "free-icon")]/@href'))
    icon_page = etree.HTML(get(icon_page_url).text)
    icon_url = json.loads(
        icon_page.xpath(
            '//script[@type="application/ld+json"][contains(text(), "cdn-icons")]/text()'
        )[0]
    )["contentUrl"]
    return Image.open(io.BytesIO(get(icon_url).content))


def get_translate(description: str, target_lang: str = "zh-CN") -> str:
    """Get a translation of the given description."""

    url = "https://translate.googleapis.com/translate_a/single"
    params = {
        "client": "gtx",
        "sl": "en",
        "tl": target_lang,
        "dt": "t",
        "q": description,
    }
    response = get(url, params=params)
    return response.json()[0][0][0]


def _make_parser() -> Callable[[str], dict]:
    """Internal function to encapsulate the parser."""

    pp.ParserElement.enable_packrat()

    chrs = "()[]{}<>«»"
    nested = functools.reduce(
        or_,
        [
            pp.nested_expr(opener, closer)
            for opener, closer in zip(chrs[::2], chrs[1::2])
        ],
    )

    sentence = pp.original_text_for(
        pp.SkipTo(
            pp.Char(".!?") + pp.Char("”’'\"")[...],
            include=True,
            ignore=nested,  # type: ignore
        )
    ).add_parse_action(lambda t: re.sub(r"\s+", " ", t[0].strip()))

    meaning = pp.Group(
        pp.Combine(ppc.integer + ".")
        + pp.QuotedString(quote_char="(", end_quote_char=")")("meaning")
        + pp.LineEnd()
        + pp.Group(
            pp.Combine(pp.Char(pp.alphas.lower()) + ".").suppress()
            + pp.SkipTo(pp.LineEnd(), include=True)("meaning")
            + pp.Group(sentence("spanish") + sentence("english"))("example")
        )[(1,)]("case")
    )[(1,)]("meaning")

    def _parse(str):
        return meaning.parseString(str).as_dict()

    return _parse


parse = _make_parser()


PLACE_HOLDER = object()


def find_dict(data: list[dict], key, val, default=PLACE_HOLDER):
    """Find the first dictionary in the given list that has the given key and value."""

    for dct in data:
        if isinstance(dct, dict) and dct.get(key, PLACE_HOLDER) == val:
            return dct
        elif (attr_dct := getattr(dct, "__dict__", None)) is not None:
            if attr_dct.get(key, PLACE_HOLDER) == val:
                return dct
        else:
            raise ValueError(f"Unable to search for {dct!r}")
    if default is not PLACE_HOLDER:
        return default
    raise ValueError(f"Unable to find {key!r}: {val!r}")


@functools.lru_cache(maxsize=1)
def get_model():
    import gensim.downloader

    return gensim.downloader.load("glove-twitter-25")


def cosine_similarity(a, b):
    return a @ b / (np.linalg.norm(a) * np.linalg.norm(b))


def pair_wise_cosine_similarity(a, b):
    return np.array(
        [[cosine_similarity(a[i], b[j]) for j in range(len(b))] for i in range(len(a))]
    )


def similar(sentence_a, sentence_b, pooling: str = "mean"):
    model = get_model()

    word_sentence_a, word_sentence_b = [
        [word for word in word_tokenize(sentence.lower()) if word in model]
        for sentence in [sentence_a, sentence_b]
    ]
    if len(word_sentence_a) == 0 or len(word_sentence_b) == 0:
        return 0

    similarity = pair_wise_cosine_similarity(
        *[
            np.array([model[word] for word in sentence])
            for sentence in [word_sentence_a, word_sentence_b]
        ]  # convert to embedding
    )

    if pooling == "mean":
        return similarity.mean()
    elif pooling == "max":
        return similarity.max()
    elif pooling == "min":
        return similarity.min()
    else:
        raise ValueError(f"Unknown pooling method: {pooling!r}")


def get_example_sentence(word: str, definition: str):
    try:
        resp = get(f"https://www.spanishdict.com/translate/{word}")
        resp.raise_for_status()
    except requests.exceptions.HTTPError as e:
        raise ValueError(f"Unable to get example sentence for {word!r}") from e

    html = etree.HTML(resp.text)
    text = ""

    for element in html.xpath('//*[@id="dictionary-neodict-es"]/div/div/div'):
        temp = "\n".join(pq(e).text() for e in element.xpath(f"./div/div"))
        if temp.startswith("A"):
            text += "\n"
            continue
        text += temp

    meanings = parse(text)["meaning"]

    top_meanings = [
        item["meaning"] + " " + " ".join(case["meaning"] for case in item["case"])
        for item in meanings
    ]

    similarities = [similar(definition, meaning, "max") for meaning in top_meanings]
    idx = np.argmax(similarities)
    sub_meanings = meanings[idx]["case"]
    examples = [case["example"] for case in sub_meanings]

    similarities = [
        similar(definition, sentence["english"], "max") for sentence in examples
    ]
    idx = np.argmax(similarities)
    sentence = examples[idx]["spanish"]
    return sentence


from functools import cached_property


class Vocab:
    def __init__(
        self,
        word: str,
        definition: str,
        image_style: Literal["image", "icon"] = "image",
        color_style: Literal["color", "gradient", "black"] = "color",
        shape_style: Literal["fill", "outline", "hand-drawn"] = "fill",
    ):
        self.word = word
        self.definition = definition
        self._word = self.normalize_word(word)

        self.image_style = image_style
        if color_style not in ["color", "gradient", "black"]:
            raise ValueError(f"Unknown color style: {color_style!r}")
        self.color_style = color_style
        if shape_style not in ["fill", "outline", "hand-drawn"]:
            raise ValueError(f"Unknown shape style: {shape_style!r}")
        self.shape_style = shape_style

    @cached_property
    def example(self):
        try:
            return get_example_sentence(self._word, self.definition)
        except Exception as e:
            print(repr(e))
            logger.error(f"Unable to get example sentence for {self.word!r}")
            words = self._word.split()
            if len(words) > 1:
                lo, hi = 0, 0
                results = {}
                while hi < len(words):
                    try:
                        hi += 1
                        sub_word = " ".join(words[lo:hi])
                        logger.info(f"Retrying with {sub_word!r}")
                        results[sub_word] = get_example_sentence(
                            sub_word, self.definition
                        )
                    except Exception:
                        logger.error(f"Retrying with {sub_word!r} failed")
                    else:
                        lo += 1

                if len(results) > 0:
                    logger.info(
                        f"Successfully retrieved example sentences for {self.word!r}"
                    )
                    return results[max(results, key=lambda k: len(k))]

    @cached_property
    def image(self):
        try:
            if self.image_style == "image":
                return get_image(self._word)
            elif self.image_style == "icon":
                try:
                    return get_icon(self._word, self.color_style, self.shape_style)
                except Exception:
                    logger.error(f"Unable to get icon for {self.word!r}.")

                    for word in sorted(self.description.split(), key=len, reverse=True):
                        try:
                            return get_icon(word, self.color_style, self.shape_style)
                        except Exception:
                            pass
                    for word in self._word.split():
                        try:
                            return get_icon(word, self.color_style, self.shape_style)
                        except Exception:
                            pass
                    raise ValueError(f"Retry failed for {self.word!r}")
        except Exception:
            logger.error(f"Unable to get image for {self.word!r}")
            return None

    @staticmethod
    def normalize_word(word):
        if word.endswith("@"):
            word = word[:-1] + "o"  # e.g., lent@ -> lento

        # handle parenthesis
        def _handle_paren(s, start, tokens):
            content = re.split(
                r",|\s+|\s+(?:etc|and|or|y|o|e|ni)\s+", s[tokens[0] : tokens[-1]][1:-1]
            )
            content = [
                c
                for c in content
                if c.strip() and c not in {"etc", "and", "or", "y", "o", "e", "ni"}
            ]
            if len(content) > 3:
                return content[0]
            return ""

        paren = pp.original_text_for(
            pp.nestedExpr("(", ")")
            | pp.nestedExpr("[", "]")
            | pp.nestedExpr("{", "}")
            | pp.nestedExpr("<", ">")
        ).set_parse_action(_handle_paren)
        word = paren.transform_string(word).strip()

        # now handle ellipsis
        word = word.replace("...", " ")

        return word.lower()


TITLE = "Diccionario Visual"
DESCRIPTION = (
    "Escoge 18 palabras del vocabulario de esta unidad. "
    "Escribe la palabra de vocabulario y una frase completa con la palabra. Dibuja una foto que representa la palabra."
)
HEADER = (
    "Nombre: {name}",
    "Hora: 30 minutos",
)


def chunks(lst, n):
    """Yield successive n-sized chunks from lst."""
    for i in range(0, len(lst), n):
        yield lst[i : i + n]


def _visual_vocab_pipeline(cards: list[list[str]], env):
    out_dir = Path(env.get("output_dir", "out"))
    out = out_dir / "visual vocab.docx"
    style = env.get("style", "image")
    color_style = env.get("color_style", "color")
    shape_style = env.get("shape_style", "fill")

    vocabs = [
        Vocab(w, d, style, color_style, shape_style)
        for w, d in random.sample(cards, 18)
    ]

    doc = docx.Document()

    header = [line.format(name=env.get("name", "Student")) for line in HEADER]
    doc.sections[0].header.paragraphs[0].text = "\t".join(header)
    doc.add_heading(TITLE, 0)
    doc.add_paragraph(DESCRIPTION)

    def add_triplet_vocab(vocabs):
        """
        3×3 column×row table
        each column correspond to a vocab
        row 1: | Vocabulario: {word} |
        row 2: | Frase Completa: {example} |
        row 3: | Foto/Meme: {image} |
        """

        table = doc.add_table(rows=3, cols=3)
        table.style = "Table Grid"

        # add vocab
        for i, vocab in enumerate(vocabs):
            table.cell(0, i).text = f"Vocabulario: {vocab.word}"
            table.cell(1, i).text = f"Frase Completa: {vocab.example}"
            img_cell = table.cell(2, i)
            par = img_cell.paragraphs[0]
            par.text = "Foto/Meme: "

            height = docx.shared.Inches(1.5)

            if vocab.image is None:
                continue

            width = vocab.image.width / vocab.image.height * docx.shared.Inches(1.5)

            if width > docx.shared.Inches(2.5):
                width = docx.shared.Inches(2.5)
                height = (
                    vocab.image.height / vocab.image.width * docx.shared.Inches(2.5)
                )

            buf = io.BytesIO()
            vocab.image.save(buf, format="PNG")

            img_cell.add_paragraph().add_run().add_picture(
                buf,
                height=height,
                width=width,
            )

        # add some space
        doc.add_paragraph("\n\n")
        doc.add_paragraph("\n\n")

    bar = tqdm.tqdm(total=len(vocabs), leave=False)

    with ThreadPoolExecutor() as executor:
        futures = [
            executor.submit(add_triplet_vocab, vocabs) for vocabs in chunks(vocabs, 3)
        ]
        for future in as_completed(futures):
            try:
                future.result()
            except Exception as e:
                logger.error(e)
            bar.update(3)

    doc.save(out)


@register("visual-vocab")
def visual_vocab_pipeline(cards: list[list[str]], env):
    try:
        _visual_vocab_pipeline(cards, env)
    except Exception as e:
        logger.error(e)
        try:
            rm_dir(temp_dir)
        except:
            pass
